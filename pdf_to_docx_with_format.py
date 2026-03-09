#!/usr/bin/env python3
"""
PDF转Word工具（保持格式一致）
将PDF文件转换为docx文件，尝试保持格式一致
"""

import os
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def detect_heading_level(text):
    """
    检测标题级别
    
    Args:
        text: 文本内容
    
    Returns:
        标题级别（1-6，0表示不是标题）
    """
    # 标题通常较短
    if len(text) > 150:
        return 0
    
    # 标题通常以特定符号结尾
    if text.endswith((':', '：', '!', '！')):
        return 2
    
    # 标题通常包含特定关键词
    heading_keywords = ['报告', '摘要', '目录', '引言', '结论', '建议', '章节', '部分', '第一章', '第二章', '第三章', '第一节', '第二节', '第三节']
    for keyword in heading_keywords:
        if keyword in text:
            # 根据关键词确定标题级别
            if any(keyword in text for keyword in ['报告', '摘要', '目录']):
                return 1
            elif any(keyword in text for keyword in ['引言', '结论', '建议']):
                return 2
            elif any(keyword in text for keyword in ['第一章', '第二章', '第三章']):
                return 2
            elif any(keyword in text for keyword in ['第一节', '第二节', '第三节']):
                return 3
            else:
                return 2
    
    # 标题通常全大写或首字母大写
    if text.isupper():
        return 2
    elif text.istitle():
        return 3
    
    # 不是标题
    return 0


def pdf_to_docx_with_format(pdf_path, output_path):
    """
    将PDF文件转换为docx文件，尝试保持格式一致
    
    Args:
        pdf_path: PDF文件路径
        output_path: 输出docx文件路径
    """
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体和字号
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # 使用宋体
    font.size = Pt(12)  # 设置字号为12
    
    # 创建标题样式
    for i in range(1, 4):
        heading_style = doc.styles.add_style(f'CustomHeading{i}', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'SimSun'
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(16)
        elif i == 2:
            heading_font.size = Pt(14)
        else:
            heading_font.size = Pt(12)
        heading_paragraph = heading_style.paragraph_format
        heading_paragraph.space_after = Pt(12)
    
    # 读取PDF文件
    reader = PdfReader(pdf_path)
    
    print(f"处理PDF文件: {pdf_path}")
    print(f"总页数: {len(reader.pages)}")
    
    # 遍历PDF页面
    for page_num in range(len(reader.pages)):
        print(f"处理第 {page_num + 1} 页...")
        
        page = reader.pages[page_num]
        
        # 提取文本
        text = page.extract_text()
        
        if text:
            # 按行分割文本
            lines = text.split('\n')
            
            # 处理每一行
            for line in lines:
                line = line.strip()
                
                if line:
                    # 检测是否为标题
                    heading_level = detect_heading_level(line)
                    
                    if heading_level > 0:
                        # 作为标题处理
                        if heading_level == 1:
                            para = doc.add_heading(line, level=1)
                        elif heading_level == 2:
                            para = doc.add_heading(line, level=2)
                        else:
                            para = doc.add_heading(line, level=3)
                    else:
                        # 作为普通段落处理
                        para = doc.add_paragraph(line)
            
            # 除了最后一页，添加分页符
            if page_num < len(reader.pages) - 1:
                doc.add_page_break()
        else:
            print(f"警告: 第 {page_num + 1} 页没有提取到文本")
    
    # 保存Word文档
    doc.save(output_path)
    print(f"转换完成: {pdf_path} -> {output_path}")
    print(f"输出文件大小: {os.path.getsize(output_path) / 1024:.2f} KB")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="PDF转Word工具（保持格式一致）")
    parser.add_argument("pdf_path", help="PDF文件路径")
    parser.add_argument("-o", "--output", help="输出Word文件路径", default=None)
    
    args = parser.parse_args()
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        # 默认输出路径：与PDF同目录，相同文件名，.docx后缀
        base_name = os.path.splitext(args.pdf_path)[0]
        output_path = f"{base_name}_with_format.docx"
    
    # 执行转换
    pdf_to_docx_with_format(args.pdf_path, output_path)
