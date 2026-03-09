#!/usr/bin/env python3
"""
Word转PDF工具
将docx文件转换为PDF文件
"""

import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch


def docx_to_pdf(docx_path, output_path):
    """
    将Word文件转换为PDF文件
    
    Args:
        docx_path: Word文件路径
        output_path: 输出PDF文件路径
    """
    # 读取Word文件
    doc = Document(docx_path)
    
    # 创建PDF文档
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    print(f"处理Word文件: {docx_path}")
    print(f"总段落数: {len(doc.paragraphs)}")
    
    # 遍历Word文档的段落
    for i, para in enumerate(doc.paragraphs):
        if i % 10 == 0:
            print(f"处理第 {i} 个段落...")
        
        # 获取段落文本
        text = para.text.strip()
        
        if text:
            # 检查是否为标题
            is_title = False
            style = 'Normal'
            
            # 检查段落样式
            if para.style.name.startswith('Heading'):
                is_title = True
                # 根据标题级别选择样式
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                if level == 1:
                    style = 'Heading1'
                elif level == 2:
                    style = 'Heading2'
                else:
                    style = 'Heading3'
            
            # 创建段落
            p = Paragraph(text, styles[style])
            story.append(p)
            story.append(Spacer(1, 0.2*inch))
    
    # 检查是否有表格
    if len(doc.tables) > 0:
        print(f"发现 {len(doc.tables)} 个表格")
        # 简单处理表格，只添加表格存在的提示
        story.append(Paragraph("[表格内容]", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # 构建PDF
    print("构建PDF文件...")
    pdf.build(story)
    
    print(f"转换完成: {docx_path} -> {output_path}")
    print(f"输出文件大小: {os.path.getsize(output_path) / 1024:.2f} KB")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Word转PDF工具")
    parser.add_argument("docx_path", help="Word文件路径")
    parser.add_argument("-o", "--output", help="输出PDF文件路径", default=None)
    
    args = parser.parse_args()
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        # 默认输出路径：与Word同目录，相同文件名，.pdf后缀
        base_name = os.path.splitext(args.docx_path)[0]
        output_path = f"{base_name}.pdf"
    
    # 执行转换
    docx_to_pdf(args.docx_path, output_path)
