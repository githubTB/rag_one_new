#!/usr/bin/env python3
"""
重新解析PDF工具
使用pdf技能目录下的工具重新解析PDF文件
"""

import os
import sys
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def reparse_pdf(pdf_path, output_path):
    """
    重新解析PDF文件，创建一个更完整的Word文档
    
    Args:
        pdf_path: PDF文件路径
        output_path: 输出Word文件路径
    """
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体和字号
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # 使用宋体
    font.size = Pt(12)  # 设置字号为12
    
    # 读取PDF文件
    reader = PdfReader(pdf_path)
    
    print(f"重新解析PDF文件: {pdf_path}")
    print(f"总页数: {len(reader.pages)}")
    
    # 遍历PDF页面
    for page_num in range(len(reader.pages)):
        print(f"处理第 {page_num + 1} 页...")
        
        page = reader.pages[page_num]
        
        # 提取文本
        text = page.extract_text()
        
        if text:
            # 按行分割文本
            lines = text.split('\n')
            
            # 处理每一行
            for line in lines:
                line = line.strip()
                
                if line:
                    # 检测是否为标题
                    is_title = False
                    
                    # 标题通常较短
                    if len(line) < 150:
                        # 标题通常全大写或首字母大写
                        if line.isupper() or line.istitle():
                            is_title = True
                        # 标题通常以特定符号结尾
                        elif line.endswith((':', '：', '!', '！')):
                            is_title = True
                        # 标题通常包含特定关键词
                        elif any(keyword in line for keyword in ['报告', '摘要', '目录', '引言', '结论', '建议', '章节', '部分']):
                            is_title = True
                    
                    if is_title:
                        # 作为标题处理
                        para = doc.add_heading(line, level=2)
                    else:
                        # 作为普通段落处理
                        para = doc.add_paragraph(line)
            
            # 除了最后一页，添加分页符
            if page_num < len(reader.pages) - 1:
                doc.add_page_break()
        else:
            print(f"警告: 第 {page_num + 1} 页没有提取到文本")
    
    # 保存Word文档
    doc.save(output_path)
    print(f"重新解析完成: {pdf_path} -> {output_path}")
    print(f"输出文件大小: {os.path.getsize(output_path) / 1024:.2f} KB")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="重新解析PDF工具")
    parser.add_argument("pdf_path", help="PDF文件路径")
    parser.add_argument("-o", "--output", help="输出Word文件路径", default=None)
    
    args = parser.parse_args()
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        # 默认输出路径：与PDF同目录，相同文件名，.docx后缀
        base_name = os.path.splitext(args.pdf_path)[0]
        output_path = f"{base_name}_reparsed.docx"
    
    # 执行重新解析
    reparse_pdf(args.pdf_path, output_path)
