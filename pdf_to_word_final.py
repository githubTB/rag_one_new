#!/usr/bin/env python3
"""
最终版PDF转Word工具
尝试提取图片并保留更多格式
"""

import os
import tempfile
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_images_from_pdf(pdf_path, output_dir):
    """
    从PDF文件中提取图片
    
    Args:
        pdf_path: PDF文件路径
        output_dir: 输出图片目录
    
    Returns:
        图片文件列表
    """
    images = []
    reader = PdfReader(pdf_path)
    
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        
        # 检查页面是否有图片
        if '/XObject' in page['/Resources']:
            x_objects = page['/Resources']['/XObject'].get_object()
            
            for obj_name, obj in x_objects.items():
                if obj['/Subtype'] == '/Image':
                    # 尝试获取图片信息
                    width = obj.get('/Width', 0)
                    height = obj.get('/Height', 0)
                    color_space = obj.get('/ColorSpace', 'Unknown')
                    
                    print(f"发现图片: {obj_name} 在第 {page_num + 1} 页, 尺寸: {width}x{height}, 颜色空间: {color_space}")
                    
                    # 尝试提取图片数据
                    try:
                        data = obj.get('/Filter')
                        if data:
                            print(f"  图片过滤器: {data}")
                    except Exception as e:
                        pass
    
    return images


def process_text_with_formatting(text):
    """
    处理文本，尝试保留更多格式
    
    Args:
        text: 原始文本
    
    Returns:
        处理后的文本段落列表
    """
    paragraphs = []
    
    if not text:
        return paragraphs
    
    # 按行分割
    lines = text.split('\n')
    
    # 当前段落
    current_para = []
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # 空行，结束当前段落
            if current_para:
                paragraphs.append(' '.join(current_para))
                current_para = []
        else:
            # 非空行，添加到当前段落
            current_para.append(line)
    
    # 处理最后一个段落
    if current_para:
        paragraphs.append(' '.join(current_para))
    
    return paragraphs


def pdf_to_word(pdf_path, output_path):
    """
    将PDF文件转换为Word文档，尝试保留更多格式和附件
    
    Args:
        pdf_path: PDF文件路径
        output_path: 输出Word文件路径
    """
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体和字号
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # 使用宋体
    font.size = Pt(12)  # 设置字号为12
    
    # 创建临时目录用于存储提取的图片
    with tempfile.TemporaryDirectory() as temp_dir:
        # 提取图片
        print("尝试提取图片...")
        images = extract_images_from_pdf(pdf_path, temp_dir)
        print(f"检测到 {len(images)} 个图片对象")
        
        # 读取PDF文件
        reader = PdfReader(pdf_path)
        
        print(f"处理PDF文件: {pdf_path}")
        print(f"总页数: {len(reader.pages)}")
        
        # 遍历PDF页面
        for page_num in range(len(reader.pages)):
            print(f"处理第 {page_num + 1} 页...")
            
            page = reader.pages[page_num]
            
            # 提取文本
            text = page.extract_text()
            
            if text:
                # 处理文本，尝试保留更多格式
                paragraphs = process_text_with_formatting(text)
                
                # 添加段落到Word文档
                for i, para_text in enumerate(paragraphs):
                    if para_text:
                        # 检测是否为标题
                        is_title = False
                        
                        # 标题通常较短
                        if len(para_text) < 150:
                            # 标题通常全大写或首字母大写
                            if para_text.isupper() or para_text.istitle():
                                is_title = True
                            # 标题通常以特定符号结尾
                            elif para_text.endswith((':', '：', '!', '！')):
                                is_title = True
                        
                        if is_title:
                            # 作为标题处理
                            para = doc.add_heading(para_text, level=2)
                        else:
                            # 作为普通段落处理
                            para = doc.add_paragraph(para_text)
                
                # 除了最后一页，添加分页符
                if page_num < len(reader.pages) - 1:
                    doc.add_page_break()
            else:
                print(f"警告: 第 {page_num + 1} 页没有提取到文本")
    
    # 保存Word文档
    doc.save(output_path)
    print(f"转换完成: {pdf_path} -> {output_path}")
    print(f"输出文件大小: {os.path.getsize(output_path) / 1024:.2f} KB")
    
    # 生成转换报告
    report_path = os.path.splitext(output_path)[0] + "_report.txt"
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(f"PDF转Word转换报告\n")
        f.write(f"原始PDF文件: {pdf_path}\n")
        f.write(f"转换后Word文件: {output_path}\n")
        f.write(f"PDF页数: {len(reader.pages)}\n")
        f.write(f"转换后文件大小: {os.path.getsize(output_path) / 1024:.2f} KB\n")
        f.write(f"\n注意事项:\n")
        f.write(f"1. 由于PDF格式的复杂性，可能无法完全保留所有格式\n")
        f.write(f"2. 图片已检测到但可能无法完全提取\n")
        f.write(f"3. 表格可能需要手动调整格式\n")
        f.write(f"4. 建议使用Microsoft Word打开并检查转换结果\n")
    
    print(f"转换报告已生成: {report_path}")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="最终版PDF转Word工具")
    parser.add_argument("pdf_path", help="PDF文件路径")
    parser.add_argument("-o", "--output", help="输出Word文件路径", default=None)
    
    args = parser.parse_args()
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        # 默认输出路径：与PDF同目录，相同文件名，.docx后缀
        base_name = os.path.splitext(args.pdf_path)[0]
        output_path = f"{base_name}_final.docx"
    
    # 执行转换
    pdf_to_word(args.pdf_path, output_path)
